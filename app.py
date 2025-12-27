import streamlit as st
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Flowable, Table, TableStyle, PageBreak, Image
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
import io
from datetime import date
import re
try:
    from PyPDF2 import PdfReader
except ImportError:
    try:
        from pypdf import PdfReader
    except ImportError:
        PdfReader = None

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(page_title="Gerador de Currículos ATS", page_icon="favicon.png", layout="wide")

# --- CONSTANTES & I18N (Internationalization) ---
BLUE_COLOR = "#0056b3"

TRANSLATIONS = {
    'pt': {
        'app_title': "Gerador de Currículos & Cover Letter (ATS-Friendly)",
        'tab_resume': "Currículo",
        'tab_cover': "Cover Letter",
        'download_btn': "📥 Baixar Currículo em PDF",
        'download_cl_btn': "📥 Baixar Carta em PDF",
        'preview_title': "👁️ Pré-visualização do Documento",
        'settings': "⚙️ Configurações",
        'density_label': "Densidade do Layout",
        'density_help': "Ajuste a densidade para caber mais conteúdo.",
        'lang_label': "Idioma / Language",
        'section_editor': "📝 Editor de Seções",
        'go_to': "Ir para:",
        'tab_proposal': "Proposta Comercial (ABNT)",
        'tab_report': "Relatório Técnico (ABNT)",
        'prop_header': "Gerador de Propostas (Padrão PMI/ABNT)",
        'rep_header': "Gerador de Relatórios (Padrão ABNT)",
        'lbl_institution': "Nome da Instituição (Opcional)",
        'lbl_author_prop': "Nome do Autor/Responsável",
        'lbl_subtitle': "Subtítulo",
        'lbl_city': "Cidade",
        'lbl_year': "Ano",
        'lbl_theme': "Tema/Natureza do Trabalho (Folha de Rosto)",
        'lbl_add_sub': "Adicionar Subseção",
        'lbl_sub_title': "Título da Subseção",
        'lbl_sub_content': "Conteúdo",
        'lbl_client_name': "Nome do Cliente",
        'lbl_project_title': "Título do Projeto/Relatório",
        'lbl_date': "Data",
        'lbl_milestone': "Marco/Entrega",
        'lbl_budget_item': "Descrição do Item",
        'lbl_amount': "Valor",
        'lbl_img_upload': "Upload de Imagem (Anexo)",
        'lbl_img_caption': "Legenda da Imagem",
        'download_prop_btn': "📥 Baixar Proposta em PDF",
        'download_rep_btn': "📥 Baixar Relatório em PDF",
        # Section Titles (Proposal)
        'sec_1': "Resumo Executivo",
        'sec_2': "Declaração do Problema",
        'sec_3': "Objetivos do Projeto",
        'sec_4': "Estratégia de Implementação",
        'sec_5': "Cronograma do Projeto",
        'sec_6': "Orçamento",
        'sec_7': "Entregáveis",
        'sec_8': "Gestão de Riscos",
        'sec_9': "Conclusão",
        # Section Titles (Report)
        'rep_sec_sum': "Sumário",
        'rep_sec_1': "Introdução",
        'rep_sec_2': "Objetivo",
        'rep_sec_3': "Material ou Métodos",
        'rep_sec_4': "Desenvolvimento",
        'rep_sec_5': "Resultados",
        'rep_sec_6': "Conclusão",
        'rep_sec_7': "Referências",
        'rep_sec_8': "Anexos",
        # Resume Headers
        'contact_header': "Informações de Contato",
        'summary_header': "RESUMO EXECUTIVO",
        'skills_header': "COMPETÊNCIAS-CHAVE",
        'experience_header': "EXPERIÊNCIA PROFISSIONAL",
        'education_header': "EDUCAÇÃO",
        'certifications_header': "CERTIFICAÇÕES E CURSOS",
        'languages_header': "IDIOMAS",
        'awards_header': "RECONHECIMENTO E PRÊMIOS",
        'volunteering_header': "VOLUNTARIADO",
        # Cover Letter Headers/Labels
        'cl_recipient_header': "Destinatário",
        'cl_hook_header': "Abertura & Hook",
        'cl_narrative_header': "Narrativa & Contexto",
        'cl_competencies_header': "Competencies (STAR)",
        'cl_alignment_header': "Alinhamento & Diferenciais",
        'cl_closing_header': "Fechamento",
        'lbl_manager': "Nome do Gerente / Recrutador",
        'lbl_company_cl': "Nome da Empresa",
        'lbl_address': "Endereço da Empresa (Opcional)",
        'lbl_greeting': "Saudação (Ex: Prezado Sr. Silva)",
        'lbl_hook': "Frase de Impacto (Hook) - Por que você?",
        'lbl_narrative': "Sua História / Filosofia Profissional",
        'lbl_star1': "Competência 1 (Situação-Ação-Resultado)",
        'lbl_star2': "Competência 2 (Situação-Ação-Resultado)",
        'lbl_alignment': "Pesquisa sobre a Empresa & Conexão",
        'lbl_differentiation': "Seus Diferenciais Únicos (Lista)",
        'lbl_closing': "Chamada para Ação & Despedida",
        # Form Labels Resume
        'lbl_name': "Nome Completo",
        'lbl_email': "Email",
        'lbl_phone': "Telefone",
        'lbl_linkedin': "LinkedIn / Portfólio",
        'lbl_location': "Localização (Cidade, Estado)",
        'lbl_summary': "Descreva seu perfil profissional",
        'lbl_skills': "Liste suas habilidades separadas por vírgula",
        'lbl_company': "Empresa",
        'lbl_position': "Cargo/Posição",
        'lbl_start': "Início",
        'lbl_end': "Fim",
        'lbl_desc': "Descrição (Use 'Enter' para criar novos parágrafos)",
        'lbl_institution': "Instituição",
        'lbl_degree': "Grau / Curso",
        'lbl_year': "Ano",
        'lbl_issuer': "Emissor",
        'lbl_updated': "Data de Atualização",
        'footer_updated': "Atualizado em",
        # Languages Labels
        'lbl_language': "Idioma",
        'lbl_conv': "Conversação",
        'lbl_comp': "Compreensão",
        'lbl_writ': "Escrita",
        'lbl_level_basic': "Básico",
        'lbl_level_inter': "Intermediário",
        'lbl_level_adv': "Avançado",
        'lbl_level_fluent': "Fluente",
        'lbl_level_native': "Nativo",
        # Awards & Volunteering Labels
        'lbl_award_title': "Nome do Prêmio/Evento",
        'lbl_award_issuer': "Órgão Emissor",
        'lbl_award_date': "Data do Prêmio",
        'lbl_vol_role': "Papel/Função",
        'lbl_vol_org': "Organização",
        'lbl_vol_cat': "Categoria (ex: Social)",
        'connector_offered_by': "Oferecido por",
        'btn_add': "Adicionar",
        'btn_update': "Atualizar",
        'btn_save': "Salvar",
        'btn_save': "Salvar",
        'btn_remove': "Remover",
        # ATS Simulator
        'tab_ats': "Simulador ATS",
        'ats_header': "Simulador de Leitura ATS (Algoritmo Puro)",
        'ats_desc': "Este módulo simula como um robô (ATS) lê seu currículo. Ele extrai o texto bruto e tenta identificar seções sem usar IA, apenas padrões.",
        'ats_upload': "Faça upload do seu PDF para análise",
        'ats_score': "Pontuação de Legibilidade ATS",
        'ats_text_len': "Caracteres Lidos",
        'ats_sec_found': "Seções Identificadas",
        'ats_raw_text': "Texto Bruto Extraído (O que o robô vê)",
        'ats_parsed_data': "Dados Estruturados Identificados",
        'ats_error_lib': "A biblioteca PyPDF2 não está instalada. Instale com 'pip install PyPDF2' para usar este módulo."
    },
    'en': {
        'app_title': "Resume & Cover Letter Builder (ATS-Friendly)",
        'tab_resume': "Resume",
        'tab_cover': "Cover Letter",
        'download_btn': "📥 Download Resume as PDF",
        'download_cl_btn': "📥 Download Cover Letter as PDF",
        'preview_title': "👁️ Document Preview",
        'settings': "⚙️ Settings",
        'density_label': "Layout Density",
        'density_help': "Adjust density to fit more content.",
        'lang_label': "Language / Idioma",
        'section_editor': "📝 Section Editor",
        'go_to': "Go to:",
        'tab_proposal': "Proposal (ABNT/PMI)",
        'prop_header': "Proposal Builder (Professional)",
        'lbl_institution': "Institution Name (Optional)",
        'lbl_author_prop': "Author Name",
        'lbl_subtitle': "Project Subtitle",
        'lbl_city': "City",
        'lbl_year': "Year",
        'lbl_theme': "Theme/Context (Title Page)",
        'lbl_add_sub': "Add Subsection",
        'lbl_sub_title': "Subsection Title",
        'go_to': "Go to:",
        'tab_proposal': "Proposal (ABNT/PMI)",
        'tab_report': "Technical Report (ABNT)",
        'prop_header': "Proposal Builder (Professional)",
        'rep_header': "Report Builder (ABNT Standard)",
        'lbl_institution': "Institution Name (Optional)",
        'lbl_author_prop': "Author Name",
        'lbl_subtitle': "Subtitle",
        'lbl_city': "City",
        'lbl_year': "Year",
        'lbl_theme': "Theme/Context (Title Page)",
        'lbl_add_sub': "Add Subsection",
        'lbl_sub_title': "Subsection Title",
        'lbl_sub_content': "Content",
        'lbl_client_name': "Client Name",
        'lbl_project_title': "Project/Report Title",
        'lbl_date': "Date",
        'lbl_milestone': "Milestone/Deliverable",
        'lbl_budget_item': "Item Description",
        'lbl_amount': "Amount",
        'lbl_img_upload': "Upload Image (Annex)",
        'lbl_img_caption': "Image Caption",
        'download_prop_btn': "📥 Download Proposal as PDF",
        'download_rep_btn': "📥 Download Report as PDF",
        # Section Titles (Proposal)
        'sec_1': "Executive Summary",
        'sec_2': "Problem Statement",
        'sec_3': "Project Objectives",
        'sec_4': "Implementation Strategy",
        'sec_5': "Project Timeline",
        'sec_6': "Budget",
        'sec_7': "Deliverables",
        'sec_8': "Risk Management",
        'sec_9': "Conclusion",
        # Section Titles (Report)
        'rep_sec_sum': "Table of Contents",
        'rep_sec_1': "Introduction",
        'rep_sec_2': "Objective",
        'rep_sec_3': "Materials or Methods",
        'rep_sec_4': "Development",
        'rep_sec_5': "Results",
        'rep_sec_6': "Conclusion",
        'rep_sec_7': "References",
        'rep_sec_8': "Annexes",
        # Resume Headers
        'contact_header': "Contact Information",
        'summary_header': "PROFESSIONAL SUMMARY",
        'skills_header': "KEY SKILLS",
        'experience_header': "WORK EXPERIENCE",
        'education_header': "EDUCATION",
        'certifications_header': "CERTIFICATIONS",
        'languages_header': "LANGUAGES",
        'awards_header': "AWARDS & RECOGNITION",
        'volunteering_header': "VOLUNTEERING",
        # Cover Letter Headers/Labels
        'cl_recipient_header': "Recipient Info",
        'cl_hook_header': "Opening & Hook",
        'cl_narrative_header': "Narrative & Context",
        'cl_competencies_header': "Competencies (STAR)",
        'cl_alignment_header': "Alignment & Differentiation",
        'cl_closing_header': "Closing",
        'lbl_manager': "Hiring Manager Name",
        'lbl_company_cl': "Company Name",
        'lbl_address': "Company Address (Optional)",
        'lbl_greeting': "Greeting (e.g., Dear Mr. Smith)",
        'lbl_hook': "Value Proposition Hook - Why you?",
        'lbl_narrative': "Your Story / Professional Philosophy",
        'lbl_star1': "Competency 1 (Situation-Action-Result)",
        'lbl_star2': "Competency 2 (Situation-Action-Result)",
        'lbl_alignment': "Company Research & Connection",
        'lbl_differentiation': "Unique Differentiators (List)",
        'lbl_closing': "Call to Action & Sign-off",
        # Form Labels Resume
        'lbl_name': "Full Name",
        'lbl_email': "Email",
        'lbl_phone': "Phone",
        'lbl_linkedin': "LinkedIn / Portfolio",
        'lbl_location': "Location (City, State/Country)",
        'lbl_summary': "Describe your professional profile",
        'lbl_skills': "List your skills separated by commas",
        'lbl_company': "Company",
        'lbl_position': "Position/Title",
        'lbl_start': "Start Date",
        'lbl_end': "End Date",
        'lbl_desc': "Description (Use 'Enter' to create new paragraphs)",
        'lbl_institution': "Institution",
        'lbl_degree': "Degree / Course",
        'lbl_year': "Year",
        'lbl_issuer': "Issuer",
        'lbl_updated': "Last Updated Date",
        'footer_updated': "Last updated",
        # Languages Labels
        'lbl_language': "Language",
        'lbl_conv': "Conversation",
        'lbl_comp': "Comprehension",
        'lbl_writ': "Writing",
        'lbl_level_basic': "Basic",
        'lbl_level_inter': "Intermediate",
        'lbl_level_adv': "Advanced",
        'lbl_level_fluent': "Fluent",
        'lbl_level_native': "Native",
        # Awards & Volunteering Labels
        'lbl_award_title': "Award/Event Name",
        'lbl_award_issuer': "Issuer Organization",
        'lbl_award_date': "Award Date",
        'lbl_vol_role': "Role/Function",
        'lbl_vol_org': "Organization",
        'lbl_vol_cat': "Category (e.g., Social)",
        'connector_offered_by': "Offered by",
        'btn_add': "Add",
        'btn_update': "Update",
        'btn_save': "Save",
        'btn_save': "Save",
        'btn_remove': "Remove",
        # ATS Simulator
        'tab_ats': "ATS Simulator",
        'ats_header': "ATS Parsing Simulator (Pure Algorithm)",
        'ats_desc': "This module simulates how a robot (ATS) reads your resume. It extracts raw text and attempts to identify sections without AI.",
        'ats_upload': "Upload your PDF for analysis",
        'ats_score': "ATS Readability Score",
        'ats_text_len': "Characters Read",
        'ats_sec_found': "Sections Identified",
        'ats_raw_text': "Extracted Raw Text (What the robot sees)",
        'ats_parsed_data': "Identified Structured Data",
        'ats_error_lib': "PyPDF2 library is not installed. Install with 'pip install PyPDF2' to use this module."
    }
}

# --- GERENCIAMENTO DE ESTADO (SESSION STATE) ---
if 'resume_data' not in st.session_state:
    st.session_state['resume_data'] = {
        "contact": {
            "name": "Seu Nome / Your Name",
            "email": "email@example.com",
            "phone": "+55 11 99999-9999",
            "linkedin": "linkedin.com/in/profile",
            "location": "City, Country"
        },
        "summary": "Experienced professional focused on results...",
        "achievements": "",
        "stack": "Python Fullstack Senior Developer",
        "skills": ["Python", "Streamlit", "Project Management"],
        "experience": [],
        "education": [],
        "certifications": [],
        "languages": [],
        "awards": [],
        "volunteering": [],
        "updated_at": date.today()
    }

if 'cover_letter_data' not in st.session_state:
    st.session_state['cover_letter_data'] = {
        "recipient": {"manager": "", "company": "", "address": ""},
        "opening": {"greeting": "", "hook": ""},
        "narrative": "",
        "competencies": {"star1": "", "star2": ""},
        "alignment": {"research": "", "differentiation": ""},
        "closing": ""
    }

if 'proposal_data' not in st.session_state:
    st.session_state['proposal_data'] = {
        "cover": {
            "institution": "", "author": "", "title": "", "subtitle": "",
            "city": "", "year": str(date.today().year), "theme": ""
        },
        # Sections 1, 2, 3, 4, 7, 8, 9 are text-based with subsections
        "text_sections": {
            "1": [], "2": [], "3": [], "4": [], "7": [], "8": [], "9": []
        },
        # Sections 5 and 6 are special (Timeline, Budget)
        "timeline": [],
        "budget": []
    }

if 'report_data' not in st.session_state:
    st.session_state['report_data'] = {
        "cover": {
            "institution": "", "author": "", "title": "", "subtitle": "",
            "city": "", "year": str(date.today().year), "theme": ""
        },
        # Sections 1-7 are text-based
        "text_sections": {
            "1": [], "2": [], "3": [], "4": [], "5": [], "6": [], "7": []
        },
        "annexes": [] # List of {'image': bytes, 'caption': str}
    }

if 'ats_data' not in st.session_state:
    st.session_state['ats_data'] = {
        "score": 0,
        "raw_text": "",
        "sections_found": [],
        "parsed_content": {}
    }

# --- FUNÇÕES DE LÓGICA (CRUD) ---
def add_experience(company, position, start, end, description):
    st.session_state['resume_data']['experience'].append({
        "company": company, "position": position, "start": start, "end": end, "description": description
    })
def remove_experience(index): st.session_state['resume_data']['experience'].pop(index)
def add_education(institution, degree, year):
    st.session_state['resume_data']['education'].append({"institution": institution, "degree": degree, "year": year})
def remove_education(index): st.session_state['resume_data']['education'].pop(index)
def add_certification(name, issuer, year):
    st.session_state['resume_data']['certifications'].append({"name": name, "issuer": issuer, "year": year})
def remove_certification(index): st.session_state['resume_data']['certifications'].pop(index)
def add_language(name, conv, comp, writ):
    st.session_state['resume_data']['languages'].append({"name": name, "conv": conv, "comp": comp, "writ": writ})
def remove_language(index): st.session_state['resume_data']['languages'].pop(index)
def add_award(title, issuer, date_str):
    st.session_state['resume_data']['awards'].append({"title": title, "issuer": issuer, "date": date_str})
def remove_award(index): st.session_state['resume_data']['awards'].pop(index)
def add_volunteering(role, org, start, end, category):
    st.session_state['resume_data']['volunteering'].append({"role": role, "org": org, "start": start, "end": end, "category": category})
def remove_volunteering(index): st.session_state['resume_data']['volunteering'].pop(index)

# Proposal Helpers
def add_text_subsection(sec_id, title, content):
    st.session_state['proposal_data']['text_sections'][sec_id].append({'title': title, 'content': content})
def remove_text_subsection(sec_id, index):
    st.session_state['proposal_data']['text_sections'][sec_id].pop(index)
def add_timeline_milestone(date_str, milestone):
    st.session_state['proposal_data']['timeline'].append({'date': str(date_str), 'milestone': milestone})
def remove_timeline_milestone(index): st.session_state['proposal_data']['timeline'].pop(index)
def add_budget_item(item, amount):
    st.session_state['proposal_data']['budget'].append({'item': item, 'amount': float(amount)})
def remove_budget_item(index): st.session_state['proposal_data']['budget'].pop(index)

# Report Helpers
def add_report_subsection(sec_id, title, content):
    st.session_state['report_data']['text_sections'][sec_id].append({'title': title, 'content': content})
def remove_report_subsection(sec_id, index):
    st.session_state['report_data']['text_sections'][sec_id].pop(index)
def add_report_annex(image_bytes, caption):
    st.session_state['report_data']['annexes'].append({'image': image_bytes, 'caption': caption})
def remove_report_annex(index): st.session_state['report_data']['annexes'].pop(index)

# --- AUDITORIA DE SISTEMA (HEALTH CHECK) ---
def run_system_audit():
    status = {"pdf_engine": False, "session_storage": False, "data_integrity": False}
    try:
        from reportlab.platypus import SimpleDocTemplate
        status["pdf_engine"] = True
    except ImportError: status["pdf_engine"] = False
    if 'resume_data' in st.session_state: status["session_storage"] = True
    data = st.session_state.get('resume_data', {})
    if isinstance(data, dict) and "contact" in data: status["data_integrity"] = True
    return all(status.values()), status

    return all(status.values()), status

# --- FUNÇÕES DE SIMULAÇÃO ATS (EXTRAÇÃO & ANÁLISE) ---
def extract_text_from_pdf(uploaded_file):
    if PdfReader is None: return None
    try:
        reader = PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"Error: {e}"

def analyze_ats_compatibility(text):
    if not text or len(text.strip()) < 50:
        return 0, {}, []

    score = 100
    sections = {}
    
    # 1. Normalização
    text_lower = text.lower()
    
    # 2. Padrões Regex para Seções (Multilíngue PT/EN)
    patterns = {
        'contact': r'(email|e-mail|phone|telefone|celular|linkedin|github|contact|contato)',
        'experience': r'(experience|experiência|work history|histórico profissional|employment)',
        'education': r'(education|educação|formação|academic|acadêmica|degree)',
        'skills': r'(skills|habilidades|competências|technologies|tecnologias|ferramentas)',
        'summary': r'(summary|resumo|profile|perfil|about|sobre)',
        'languages': r'(languages|idiomas|linguas)',
        'certifications': r'(certifications|certificações|cursos|courses)'
    }
    
    found_sections = []
    
    # Busca heurística simples
    for sec_name, pattern in patterns.items():
        if re.search(pattern, text_lower):
            found_sections.append(sec_name)
            sections[sec_name] = "DETECTED (Content parsing requires advanced NLP)"
        else:
            sections[sec_name] = "NOT FOUND"
            score -= 10 # Penalidade por seção faltante
            
    # 3. Penalidades de Formatação
    # Excesso de espaços em branco (simboliza formatação quebrada)
    if "  " in text: 
        score -= 5
    # Texto muito curto para um CV
    if len(text) < 500:
        score -= 20
        
    # Limites
    return max(0, score), sections, found_sections

# --- CLASSES AUXILIARES DE DESIGN ---
class HorizontalLine(Flowable):
    def __init__(self, width, color=colors.black, thickness=1):
        Flowable.__init__(self)
        self.width = width
        self.color = color
        self.thickness = thickness
    def draw(self):
        self.canv.setStrokeColor(self.color)
        self.canv.setLineWidth(self.thickness)
        self.canv.line(0, 0, self.width, 0)

# --- FUNÇÃO DE GERAÇÃO DE PDF (RESUME) ---
def generate_pdf(data, scale_factor, lang_code):
    buffer = io.BytesIO()
    t = TRANSLATIONS[lang_code]
    margin = 50 * scale_factor if scale_factor > 0.9 else 40
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=margin, leftMargin=margin, topMargin=margin, bottomMargin=margin)
    content_width = letter[0] - 2 * margin
    styles = getSampleStyleSheet()
    def scaled(size, min_size=9): return max(size * scale_factor, min_size)
    
    # Estilos (Resume) - Refatorado para Times-Roman e Centralização e Cor Preta (ATS)
    style_name = ParagraphStyle('Name', parent=styles['Heading1'], fontSize=scaled(26), leading=scaled(26) * 1.2, textColor=colors.black, fontName='Times-Bold', alignment=1, spaceAfter=scaled(4))
    
    # Novo estilo para Stack (Helvetica, Centralizado)
    style_stack = ParagraphStyle('Stack', parent=styles['Normal'], fontSize=scaled(12), leading=scaled(12) * 1.2, textColor=colors.black, fontName='Helvetica', alignment=1, spaceAfter=scaled(12))

    style_contact = ParagraphStyle('Contact', parent=styles['Normal'], fontSize=scaled(10), leading=scaled(10) * 1.2, textColor=colors.black, fontName='Times-Roman', alignment=1, spaceAfter=scaled(4))
    
    # Section Header - Times-Bold e Centralizado - Cor Preta - Espaçamento Reduzido
    style_section_header = ParagraphStyle('SectionHeader', parent=styles['Heading2'], fontSize=scaled(13), leading=scaled(13) * 1.05, textColor=colors.black, fontName='Times-Bold', textTransform='uppercase', spaceBefore=scaled(12), spaceAfter=scaled(0), alignment=1)
    
    style_item_header = ParagraphStyle('ItemHeader', parent=styles['Normal'], fontSize=scaled(11.5), leading=scaled(11.5) * 1.2, fontName='Times-Bold', textColor=colors.black, spaceAfter=scaled(1), spaceBefore=0)
    style_item_sub = ParagraphStyle('ItemSub', parent=styles['Normal'], fontSize=scaled(10.5), leading=scaled(10.5) * 1.2, fontName='Times-Italic', textColor=colors.black, spaceAfter=scaled(2))
    style_normal = ParagraphStyle('NormalText', parent=styles['Normal'], fontSize=scaled(10.5), leading=scaled(10.5) * 1.4, alignment=TA_JUSTIFY, spaceAfter=scaled(3), fontName='Times-Roman', spaceBefore=0)
    
    story = []
    # Header
    story.append(Paragraph(data['contact']['name'], style_name))
    
    # Header de Stack (Logo após o nome)
    stack_text = data.get('stack', '')
    if stack_text:
        story.append(Paragraph(stack_text, style_stack))

    sep = " • "
    contact_parts = [data['contact']['phone'], data['contact']['email'], data['contact']['location'], data['contact']['linkedin']]
    story.append(Paragraph(sep.join([p for p in contact_parts if p]), style_contact))
    story.append(Spacer(1, scaled(4)))
    # story.append(HorizontalLine(content_width, color=colors.black, thickness=1.5))
    # story.append(Spacer(1, scaled(10)))

    def add_section_title(text):
        story.append(Paragraph(text.upper(), style_section_header))
        # story.append(HorizontalLine(content_width, color=colors.black, thickness=0.5))
        # Spacer removed for Modern/Minimalist look

    if data['summary']:
        add_section_title(t['summary_header'])
        for line in data['summary'].split('\n'):
            if line.strip(): story.append(Paragraph(line, style_normal))
        
        # Subseção 'Feitos' (Achievements) dentro do Summary
        achievements = data.get('achievements', '')
        if achievements:
            ach_label = "FEITOS" if lang_code == 'pt' else "KEY ACHIEVEMENTS"
            story.append(Spacer(1, scaled(4)))
            story.append(Paragraph(f"<b>{ach_label}</b>", style_normal))
            for line in achievements.split('\n'):
                if line.strip(): story.append(Paragraph(line, style_normal))

    if data['skills']:
        add_section_title(t['skills_header'])
        story.append(Paragraph(", ".join(data['skills']), style_normal))
    if data['experience']:
        add_section_title(t['experience_header'])
        for exp in data['experience']:
            story.append(Paragraph(f"{exp['position']} | {exp['company']}", style_item_header))
            story.append(Paragraph(f"{exp['start']} - {exp['end']}", style_item_sub))
            if exp['description']:
                for line in exp['description'].split('\n'):
                    if line.strip(): story.append(Paragraph(line, style_normal))
            story.append(Spacer(1, scaled(6)))
    if data['education']:
        add_section_title(t['education_header'])
        for edu in data['education']:
            story.append(Paragraph(f"{edu['degree']}", style_item_header))
            story.append(Paragraph(f"{edu['institution']} • {t['lbl_year']}: {edu['year']}", style_item_sub))
            story.append(Spacer(1, scaled(4)))
    if data['certifications']:
        add_section_title(t['certifications_header'])
        for cert in data['certifications']:
            story.append(Paragraph(f"• <b>{cert['name']}</b> ({cert['issuer']}, {cert['year']})", style_normal))
    
    # LANGUAGES SECTION
    if data.get('languages'):
        add_section_title(t['languages_header'])
        for lang in data['languages']:
            lang_text = f"• <b>{lang['name']}</b> - {t['lbl_conv']}: {lang['conv']} | {t['lbl_comp']}: {lang['comp']} | {t['lbl_writ']}: {lang['writ']}"
            story.append(Paragraph(lang_text, style_normal))

    # AWARDS SECTION
    if data.get('awards'):
        add_section_title(t['awards_header'])
        for aw in data['awards']:
            # Format: Nome | Oferecido por Órgão | Data
            aw_text = f"• {aw['title']} | {t['connector_offered_by']} {aw['issuer']} | {aw['date']}"
            story.append(Paragraph(aw_text, style_normal))

    # VOLUNTEERING SECTION
    if data.get('volunteering'):
        add_section_title(t['volunteering_header'])
        for vol in data['volunteering']:
            # Format: Nome | Org | Início -> Fim | Categoria
            vol_text = f"• <b>{vol['role']}</b> | {vol['org']} | {vol['start']} -> {vol['end']} | {vol['category']}"
            story.append(Paragraph(vol_text, style_normal))

    doc.build(story)
    return buffer.getvalue()


# --- FUNÇÃO DE GERAÇÃO DE DOCX (RESUME WORD) ---
def generate_docx(data, scale_factor, lang_code):
    if Document is None:
        return None
    
    doc = Document()
    t = TRANSLATIONS[lang_code]
    
    # Função auxiliar para escala
    def scaled(size, min_size=6): return max(size * scale_factor, min_size)

    # Configuração de Margens (Compacto/Minimalista - Escalonável)
    sections = doc.sections
    for section in sections:
        # Base margins 1.27cm (~0.5 inch) adjusted by scale
        margin_sz = 1.27 * scale_factor
        section.top_margin = Cm(margin_sz)
        section.bottom_margin = Cm(margin_sz)
        section.left_margin = Cm(margin_sz)
        section.right_margin = Cm(margin_sz)

    # Estilos Básicos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(scaled(10.5))
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0) # Minimalist - sem espaço automático
    paragraph_format.line_spacing = 1.15
    
    # Function to add formatted paragraph
    def add_para(text, bold=False, size=None, align='LEFT', space_before=0, space_after=0):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        # If size is provided, scale it; otherwise let style default handle it
        if size: run.font.size = Pt(scaled(size))
        run.font.name = 'Times New Roman'
        
        if align == 'CENTER': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'RIGHT': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Ajuste Fino de Espaçamento (Scaled)
        p.paragraph_format.space_before = Pt(scaled(space_before))
        p.paragraph_format.space_after = Pt(scaled(space_after))
        return p

    def add_section_header(text):
        # Título da Seção: Times-Bold, Caps, Sem espaço depois
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(scaled(12))
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text.upper())
        run.font.name = 'Times New Roman'
        run.bold = True
        run.font.size = Pt(scaled(13))

    # 1. Header (Nome)
    name_p = add_para(data['contact']['name'], bold=True, size=24, align='CENTER')
    
    # 2. Stack / Subtitle
    if data.get('stack'):
        stack_p = add_para(data['stack'], size=12, align='CENTER', space_after=6)
        stack_p.runs[0].font.name = 'Arial' # match PDF style

    # 3. Contact Info
    contact_parts = [data['contact']['phone'], data['contact']['email'], data['contact']['location'], data['contact']['linkedin']]
    contact_line = " • ".join([p for p in contact_parts if p])
    add_para(contact_line, size=10, align='CENTER', space_after=12)

    # 4. Summary
    if data['summary']:
        add_section_header(t['summary_header'])
        for line in data['summary'].split('\n'):
            if line.strip(): add_para(line, space_before=2)
        
        if data.get('achievements'):
             ach_label = "FEITOS" if lang_code == 'pt' else "KEY ACHIEVEMENTS"
             p = doc.add_paragraph()
             p.paragraph_format.space_before = Pt(scaled(6))
             p.paragraph_format.space_after = Pt(0)
             run = p.add_run(ach_label)
             run.bold = True
             run.font.name = 'Times New Roman'
             run.font.size = Pt(scaled(10.5))
             
             for line in data.get('achievements').split('\n'):
                 if line.strip(): add_para(line, space_before=2)

    # 5. Skills
    if data['skills']:
        add_section_header(t['skills_header'])
        add_para(", ".join(data['skills']), space_before=2)

    # 6. Experience
    if data['experience']:
        add_section_header(t['experience_header'])
        for exp in data['experience']:
            # Posição | Empresa (Bold)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(scaled(6))
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(f"{exp['position']} | {exp['company']}")
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(scaled(11.5))
            
            # Data (Italic)
            p_sub = doc.add_paragraph()
            p_sub.paragraph_format.space_after = Pt(scaled(2))
            run_sub = p_sub.add_run(f"{exp['start']} - {exp['end']}")
            run_sub.italic = True
            run_sub.font.name = 'Times New Roman'
            run_sub.font.size = Pt(scaled(10.5))

            # Descrição
            if exp['description']:
                for line in exp['description'].split('\n'):
                    if line.strip():
                        add_para(line)

    # 7. Education
    if data['education']:
        add_section_header(t['education_header'])
        for edu in data['education']:
            # Grau (Bold)
            add_para(edu['degree'], bold=True, size=11.5, space_before=6)
            # Inst - Ano (Italic)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(scaled(2))
            run = p.add_run(f"{edu['institution']} • {t['lbl_year']}: {edu['year']}")
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(scaled(10.5))

    # 8. Certifications
    if data['certifications']:
        add_section_header(t['certifications_header'])
        for cert in data['certifications']:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Cm(0.5)
            
            run_bullet = p.add_run("• ")
            run_name = p.add_run(cert['name'])
            run_name.bold = True
            run_details = p.add_run(f" ({cert['issuer']}, {cert['year']})")
            
            for r in [run_bullet, run_name, run_details]: 
                r.font.name = 'Times New Roman'
                r.font.size = Pt(scaled(10.5))

    # 9. Languages
    if data.get('languages'):
        add_section_header(t['languages_header'])
        for lang in data['languages']:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(f"• {lang['name']}")
            run.bold = True
            run2 = p.add_run(f" - {t['lbl_conv']}: {lang['conv']} | {t['lbl_comp']}: {lang['comp']} | {t['lbl_writ']}: {lang['writ']}")
            for r in [run, run2]: 
                r.font.name = 'Times New Roman'
                r.font.size = Pt(scaled(10.5))

    # 10. Awards & Volunteering (Simplificado)
    if data.get('awards'):
        add_section_header(t['awards_header'])
        for aw in data['awards']:
            add_para(f"• {aw['title']} | {aw['issuer']} | {aw['date']}")
            
    if data.get('volunteering'):
        add_section_header(t['volunteering_header'])
        for vol in data['volunteering']:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(f"• {vol['role']}")
            run.bold = True
            run2 = p.add_run(f" | {vol['org']} | {vol['category']}")
            for r in [run, run2]: r.font.name = 'Times New Roman'

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# --- FUNÇÃO DE GERAÇÃO DE PDF (COVER LETTER) ---
def generate_cl_pdf(resume_data, cl_data, scale_factor, lang_code):
    buffer = io.BytesIO()
    margin = 50 * scale_factor if scale_factor > 0.9 else 40
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=margin, leftMargin=margin, topMargin=margin, bottomMargin=margin)
    content_width = letter[0] - 2 * margin
    styles = getSampleStyleSheet()
    def scaled(size, min_size=9): return max(size * scale_factor, min_size)

    # Estilos (Cover Letter)
    style_name = ParagraphStyle('Name', parent=styles['Heading1'], fontSize=scaled(26), leading=scaled(26) * 1.2, textColor=colors.HexColor("#2C3E50"), fontName='Helvetica-Bold', alignment=1, spaceAfter=scaled(8))
    style_contact = ParagraphStyle('Contact', parent=styles['Normal'], fontSize=scaled(10), leading=scaled(10) * 1.2, textColor=colors.HexColor("#555555"), alignment=1, spaceAfter=scaled(20))
    style_body = ParagraphStyle('Body', parent=styles['Normal'], fontSize=scaled(11), leading=scaled(11) * 1.4, alignment=TA_LEFT, spaceAfter=scaled(10))
    style_recipient = ParagraphStyle('Recipient', parent=styles['Normal'], fontSize=scaled(11), leading=scaled(11) * 1.2, fontName='Helvetica-Bold', spaceAfter=scaled(20), leftIndent=0)
    
    story = []

    # 1. Header
    story.append(Paragraph(resume_data['contact']['name'], style_name))
    sep = " • "
    contact_parts = [resume_data['contact']['phone'], resume_data['contact']['email'], resume_data['contact']['location'], resume_data['contact']['linkedin']]
    story.append(Paragraph(sep.join([p for p in contact_parts if p]), style_contact))
    story.append(Spacer(1, scaled(10)))
    story.append(HorizontalLine(content_width, color=colors.HexColor(BLUE_COLOR), thickness=1.5))
    story.append(Spacer(1, scaled(20)))

    # 2. Recipient Info
    rec = cl_data['recipient']
    if rec['manager'] or rec['company']:
        rec_text = f"{date.today().strftime('%B %d, %Y')}<br/><br/>"
        if rec['manager']: rec_text += f"{rec['manager']}<br/>"
        if rec['company']: rec_text += f"{rec['company']}<br/>"
        if rec['address']: rec_text += f"{rec['address']}"
        story.append(Paragraph(rec_text, style_recipient))

    # 3. Greeting
    if cl_data['opening']['greeting']:
        story.append(Paragraph(f"{cl_data['opening']['greeting']}:", style_body))

    # 4. Hook
    if cl_data['opening']['hook']:
        story.append(Paragraph(cl_data['opening']['hook'], style_body))

    # 5. Narrative
    if cl_data['narrative']:
        for line in cl_data['narrative'].split('\n'):
            if line.strip(): story.append(Paragraph(line, style_body))

    # 6. Competencies (STAR)
    if cl_data['competencies']['star1']:
        story.append(Paragraph(cl_data['competencies']['star1'], style_body))
    if cl_data['competencies']['star2']:
        story.append(Paragraph(cl_data['competencies']['star2'], style_body))

    # 7. Alignment & Differentiation
    if cl_data['alignment']['research']:
        story.append(Paragraph(cl_data['alignment']['research'], style_body))
    if cl_data['alignment']['differentiation']:
        for line in cl_data['alignment']['differentiation'].split('\n'):
            if line.strip(): story.append(Paragraph(f"• {line}", style_body))

    # 8. Closing
    if cl_data['closing']:
        story.append(Paragraph(cl_data['closing'], style_body))
        story.append(Spacer(1, scaled(20)))
        story.append(Paragraph("Sincerely,", style_body))
        story.append(Spacer(1, scaled(30)))
        story.append(Paragraph(resume_data['contact']['name'], style_body))

    doc.build(story)
    return buffer.getvalue()

    doc.build(story)
    return buffer.getvalue()

# --- FUNÇÃO DE GERAÇÃO DE PDF (RELATÓRIO ABNT) ---
def generate_report_pdf(data, t):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=3*cm, leftMargin=3*cm, topMargin=3*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []
    
    # Estilos ABNT (Reutilizados)
    style_center = ParagraphStyle('ABNTCenter', parent=styles['Normal'], alignment=1, fontSize=12, leading=14, spaceAfter=6)
    style_title_cover = ParagraphStyle('ABNTTitle', parent=styles['Heading1'], alignment=1, fontSize=16, leading=20, fontName='Helvetica-Bold', spaceAfter=12, spaceBefore=100)
    style_subtitle = ParagraphStyle('ABNTSub', parent=styles['Normal'], alignment=1, fontSize=14, leading=16, spaceAfter=100)
    style_note = ParagraphStyle('ABNTNote', parent=styles['Normal'], alignment=TA_JUSTIFY, leftIndent=7*cm, fontSize=10, leading=12)
    
    style_h1 = ParagraphStyle('RepH1', parent=styles['Heading1'], fontSize=14, textColor=colors.black, spaceBefore=20, spaceAfter=12, textTransform='uppercase', fontName='Helvetica-Bold')
    style_h2 = ParagraphStyle('RepH2', parent=styles['Heading2'], fontSize=12, textColor=colors.black, spaceBefore=10, spaceAfter=6, fontName='Helvetica-Bold')
    style_normal = ParagraphStyle('RepNormal', parent=styles['Normal'], fontSize=11, leading=14, alignment=TA_JUSTIFY, spaceAfter=6)
    style_caption = ParagraphStyle('RepCaption', parent=styles['Normal'], fontSize=10, leading=12, alignment=1, spaceAfter=12, spaceBefore=4)

    # --- 1. CAPA ---
    if data['cover']['institution']: story.append(Paragraph(data['cover']['institution'].upper(), style_center))
    story.append(Paragraph(data['cover']['author'].upper(), style_center))
    story.append(Paragraph(data['cover']['title'].upper(), style_title_cover))
    if data['cover']['subtitle']: story.append(Paragraph(data['cover']['subtitle'], style_subtitle))
    else: story.append(Spacer(1, 100))
    story.append(Spacer(1, 200))
    story.append(Paragraph(data['cover']['city'], style_center))
    story.append(Paragraph(data['cover']['year'], style_center))
    story.append(PageBreak())

    # --- 2. FOLHA DE ROSTO ---
    story.append(Paragraph(data['cover']['author'].upper(), style_center))
    story.append(Spacer(1, 100))
    story.append(Paragraph(data['cover']['title'].upper(), style_center))
    if data['cover']['subtitle']: story.append(Paragraph(data['cover']['subtitle'], style_center))
    story.append(Spacer(1, 50))
    if data['cover']['theme']: story.append(Paragraph(data['cover']['theme'], style_note))
    story.append(Spacer(1, 200))
    story.append(Paragraph(data['cover']['city'], style_center))
    story.append(Paragraph(data['cover']['year'], style_center))
    story.append(PageBreak())

    # --- 3. SUMÁRIO (Simples) ---
    story.append(Paragraph(t['rep_sec_sum'].upper(), style_h1))
    story.append(HorizontalLine(450, color=colors.black, thickness=0.5))
    story.append(Spacer(1, 10))
    
    sections_map = [
        ("1", t['rep_sec_1']), ("2", t['rep_sec_2']), ("3", t['rep_sec_3']),
        ("4", t['rep_sec_4']), ("5", t['rep_sec_5']), ("6", t['rep_sec_6']),
        ("7", t['rep_sec_7']), ("8", t['rep_sec_8'])
    ]
    
    for sec_id, sec_title in sections_map:
        # Check if section has content
        has_content = False
        if sec_id == "8": has_content = bool(data['annexes'])
        else: has_content = bool(data['text_sections'].get(sec_id))
        
        if has_content:
            story.append(Paragraph(f"{sec_id}. {sec_title}", style_normal))
    
    story.append(PageBreak())

    # --- 4. CONTEÚDO (SEÇÕES 1-7) ---
    for sec_id, sec_title in sections_map[:-1]: # Exclude Annexes from this loop
        subsections = data['text_sections'].get(sec_id, [])
        if subsections:
            story.append(Paragraph(f"{sec_id}. {sec_title}", style_h1))
            for idx, sub in enumerate(subsections):
                if sub['title']: story.append(Paragraph(f"{sec_id}.{idx+1} {sub['title']}", style_h2))
                if sub['content']:
                    for line in sub['content'].split('\n'):
                        if line.strip(): story.append(Paragraph(line, style_normal))
            story.append(Spacer(1, 15))

    # --- 5. ANEXOS (SEÇÃO 8) ---
    if data['annexes']:
        story.append(PageBreak())
        story.append(Paragraph(f"8. {t['rep_sec_8']}", style_h1))
        for idx, annex in enumerate(data['annexes']):
            try:
                # Process Image
                img_stream = io.BytesIO(annex['image'])
                img = Image(img_stream)
                
                # Resize Logic (Max Width 15cm)
                max_width = 15 * cm
                if img.drawWidth > max_width:
                    ratio = max_width / img.drawWidth
                    img.drawWidth = max_width
                    img.drawHeight = img.drawHeight * ratio
                
                story.append(img)
                story.append(Paragraph(f"Figura {idx+1}: {annex['caption']}", style_caption))
                story.append(Spacer(1, 12))
            except Exception as e:
                story.append(Paragraph(f"[Erro ao renderizar imagem: {str(e)}]", style_normal))

    doc.build(story)
    return buffer.getvalue()

# --- FUNÇÃO DE GERAÇÃO DE PDF (PROPOSTA ABNT) ---
def generate_proposal_pdf(data, t):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=3*cm, leftMargin=3*cm, topMargin=3*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []
    
    # Estilos ABNT
    style_center = ParagraphStyle('ABNTCenter', parent=styles['Normal'], alignment=1, fontSize=12, leading=14, spaceAfter=6)
    style_title_cover = ParagraphStyle('ABNTTitle', parent=styles['Heading1'], alignment=1, fontSize=16, leading=20, fontName='Helvetica-Bold', spaceAfter=12, spaceBefore=100)
    style_subtitle = ParagraphStyle('ABNTSub', parent=styles['Normal'], alignment=1, fontSize=14, leading=16, spaceAfter=100)
    style_note = ParagraphStyle('ABNTNote', parent=styles['Normal'], alignment=TA_JUSTIFY, leftIndent=7*cm, fontSize=10, leading=12)
    
    style_h1 = ParagraphStyle('PropH1', parent=styles['Heading1'], fontSize=14, textColor=colors.HexColor(BLUE_COLOR), spaceBefore=20, spaceAfter=12, textTransform='uppercase')
    style_h2 = ParagraphStyle('PropH2', parent=styles['Heading2'], fontSize=12, textColor=colors.black, spaceBefore=10, spaceAfter=6, fontName='Helvetica-Bold')
    style_normal = ParagraphStyle('PropNormal', parent=styles['Normal'], fontSize=11, leading=14, alignment=TA_JUSTIFY, spaceAfter=6)

    # --- 1. CAPA ---
    if data['cover']['institution']:
        story.append(Paragraph(data['cover']['institution'].upper(), style_center))
    story.append(Paragraph(data['cover']['author'].upper(), style_center))
    
    story.append(Paragraph(data['cover']['title'].upper(), style_title_cover))
    if data['cover']['subtitle']:
        story.append(Paragraph(data['cover']['subtitle'], style_subtitle))
    else:
        story.append(Spacer(1, 100))
        
    story.append(Spacer(1, 200)) # Empurrar para baixo
    story.append(Paragraph(data['cover']['city'], style_center))
    story.append(Paragraph(data['cover']['year'], style_center))
    story.append(PageBreak())

    # --- 2. FOLHA DE ROSTO ---
    story.append(Paragraph(data['cover']['author'].upper(), style_center))
    story.append(Spacer(1, 100))
    story.append(Paragraph(data['cover']['title'].upper(), style_center))
    if data['cover']['subtitle']:
        story.append(Paragraph(data['cover']['subtitle'], style_center))
    
    story.append(Spacer(1, 50))
    if data['cover']['theme']:
        story.append(Paragraph(data['cover']['theme'], style_note))
    
    story.append(Spacer(1, 200)) # Empurrar para baixo
    story.append(Paragraph(data['cover']['city'], style_center))
    story.append(Paragraph(data['cover']['year'], style_center))
    story.append(PageBreak())

    # --- 3. CONTEÚDO (SEÇÕES 1-9) ---
    sections_order = [
        ("1", t['sec_1']), ("2", t['sec_2']), ("3", t['sec_3']), ("4", t['sec_4']),
        ("5", t['sec_5']), ("6", t['sec_6']), ("7", t['sec_7']), ("8", t['sec_8']), ("9", t['sec_9'])
    ]

    for sec_id, sec_title in sections_order:
        # Título da Seção (ex: 1. RESUMO EXECUTIVO)
        story.append(Paragraph(f"{sec_id}. {sec_title}", style_h1))
        story.append(HorizontalLine(450, color=colors.HexColor("#DDDDDD"), thickness=0.5))
        
        # Lógica Especial: Cronograma (5)
        if sec_id == "5": # Cronograma
            if data['timeline']:
                table_data = [[t.get('lbl_date', 'Data'), t.get('lbl_milestone', 'Marco')]]
                for item in data['timeline']: table_data.append([item['date'], item['milestone']])
                t_style = TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.HexColor(BLUE_COLOR)), ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                    ('ALIGN', (0,0), (-1,-1), 'LEFT'), ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.grey), ('BOTTOMPADDING', (0,0), (-1,0), 8), ('TOPPADDING', (0,0), (-1,0), 8)
                ])
                t_obj = Table(table_data, colWidths=[100, 350])
                t_obj.setStyle(t_style)
                story.append(t_obj)
            else: story.append(Paragraph("N/A", style_normal))

        elif sec_id == "6": # Orçamento
            if data['budget']:
                table_data = [[t.get('lbl_budget_item', 'Item'), t.get('lbl_amount', 'Valor')]]
                total = 0.0
                for item in data['budget']:
                    table_data.append([item['item'], f"{item['amount']:,.2f}"])
                    total += item['amount']
                table_data.append(['TOTAL', f"{total:,.2f}"])
                b_style = TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.HexColor(BLUE_COLOR)), ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                    ('ALIGN', (1,0), (-1,-1), 'RIGHT'), ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.grey), ('FONTNAME', (0,-1), (-1,-1), 'Helvetica-Bold'),
                    ('BACKGROUND', (0,-1), (-1,-1), colors.lightgrey), ('TEXTCOLOR', (0,-1), (-1,-1), colors.black),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 8), ('TOPPADDING', (0,0), (-1,-1), 8)
                ])
                b_obj = Table(table_data, colWidths=[350, 100])
                b_obj.setStyle(b_style)
                story.append(b_obj)
            else: story.append(Paragraph("N/A", style_normal))

        else: # Seções de Texto (1, 2, 3, 4, 7, 8, 9)
            subsections = data['text_sections'].get(sec_id, [])
            if not subsections:
                story.append(Paragraph("...", style_normal))
            for idx, sub in enumerate(subsections):
                # Subseção (ex: 1.1 Título)
                if sub['title']:
                    story.append(Paragraph(f"{sec_id}.{idx+1} {sub['title']}", style_h2))
                if sub['content']:
                    for line in sub['content'].split('\n'):
                        if line.strip(): story.append(Paragraph(line, style_normal))
        
        story.append(Spacer(1, 15))

    doc.build(story)
    return buffer.getvalue()

# --- INTERFACE DO USUÁRIO ---
def main():
    with st.sidebar:
        st.header("🌐 & ⚙️")
        lang_option = st.radio("Language / Idioma", ["Português", "English"], horizontal=True)
        lang_code = 'pt' if lang_option == "Português" else 'en'
        t = TRANSLATIONS[lang_code]

        st.divider()
        st.subheader(t['settings'])
        density_mode = st.select_slider(t['density_label'], options=["Confortável", "Normal", "Compacto", "Super Compacto"], value="Normal", help=t['density_help'])
        scale_map = {"Confortável": 1.05, "Normal": 1.0, "Compacto": 0.9, "Super Compacto": 0.85}
        scale_factor = scale_map[density_mode]
        st.session_state['resume_data']['updated_at'] = st.date_input(t['lbl_updated'], value=st.session_state['resume_data'].get('updated_at', date.today()))
        
        st.divider()
        audit_ok, audit_details = run_system_audit()
        if audit_ok: st.success("System Status: ✅ Operational")
        else: st.error(f"System Error: {audit_details}")

    st.title(t['app_title'])
    
    # --- TABS: RESUME vs COVER LETTER vs PROPOSAL vs REPORT vs ATS SIMULATOR ---
    tab1, tab2, tab3, tab4, tab5 = st.tabs([t['tab_resume'], t['tab_cover'], t.get('tab_proposal', 'Proposta'), t.get('tab_report', 'Relatório'), t.get('tab_ats', 'Simulador ATS')])

    # === TAB 1: RESUME BUILDER ===
    with tab1:
        col_edit, col_prev = st.columns([1, 1.2])
        
        with col_edit:
            st.header(t['section_editor'])
            section_map = {
                t['contact_header']: "Contact", t['summary_header']: "Summary", t['skills_header']: "Skills",
                t['experience_header']: "Experience", t['education_header']: "Education", t['certifications_header']: "Certifications",
                t['languages_header']: "Languages", t['awards_header']: "Awards", t['volunteering_header']: "Volunteering"
            }
            section_selected = st.selectbox(t['go_to'], list(section_map.keys()))
            section_logic = section_map[section_selected]

            if section_logic == "Contact":
                with st.form("contact_form"):
                    st.session_state['resume_data']['contact']['name'] = st.text_input(t['lbl_name'], st.session_state['resume_data']['contact']['name'])
                    st.session_state['resume_data']['contact']['email'] = st.text_input(t['lbl_email'], st.session_state['resume_data']['contact']['email'])
                    st.session_state['resume_data']['contact']['phone'] = st.text_input(t['lbl_phone'], st.session_state['resume_data']['contact']['phone'])
                    st.session_state['resume_data']['contact']['linkedin'] = st.text_input(t['lbl_linkedin'], st.session_state['resume_data']['contact']['linkedin'])
                    st.session_state['resume_data']['contact']['location'] = st.text_input(t['lbl_location'], st.session_state['resume_data']['contact']['location'])
                    st.session_state['resume_data']['stack'] = st.text_input("Stack / Título Profissional (Abaixo do Nome)", st.session_state['resume_data'].get('stack', ''))
                    st.form_submit_button(t['btn_update'])
            
            elif section_logic == "Summary":
                with st.form("summary_form"):
                    new_summary = st.text_area(t['lbl_summary'], st.session_state['resume_data']['summary'], height=150)
                    new_achievements = st.text_area("Principais Feitos / Achievements (Opcional - aparecerá dentro do Resumo)", st.session_state['resume_data'].get('achievements', ''), height=100)
                    if st.form_submit_button(t['btn_save']):
                        st.session_state['resume_data']['summary'] = new_summary
                        st.session_state['resume_data']['achievements'] = new_achievements
                        st.success("OK!")

            elif section_logic == "Skills":
                current_skills = ", ".join(st.session_state['resume_data']['skills'])
                with st.form("skills_form"):
                    skills_input = st.text_area(t['lbl_skills'], current_skills)
                    if st.form_submit_button(t['btn_save']):
                        st.session_state['resume_data']['skills'] = [s.strip() for s in skills_input.split(",") if s.strip()]
                        st.success("OK!")

            elif section_logic == "Experience":
                for i, exp in enumerate(st.session_state['resume_data']['experience']):
                    with st.expander(f"{exp['position']} - {exp['company']}"):
                        st.write(f"{exp['start']} - {exp['end']}")
                        if st.button(f"{t['btn_remove']} {i}", key=f"del_exp_{i}"): remove_experience(i); st.rerun()
                with st.form("add_exp_form"):
                    company = st.text_input(t['lbl_company'])
                    position = st.text_input(t['lbl_position'])
                    col1, col2 = st.columns(2)
                    start = col1.text_input(t['lbl_start'])
                    end = col2.text_input(t['lbl_end'])
                    desc = st.text_area(t['lbl_desc'], height=100)
                    if st.form_submit_button(t['btn_add']):
                        if company and position: add_experience(company, position, start, end, desc); st.rerun()

            elif section_logic == "Education":
                for i, edu in enumerate(st.session_state['resume_data']['education']):
                    with st.expander(f"{edu['degree']} - {edu['institution']}"):
                        if st.button(f"{t['btn_remove']} {i}", key=f"del_edu_{i}"): remove_education(i); st.rerun()
                with st.form("add_edu_form"):
                    inst = st.text_input(t['lbl_institution'])
                    degree = st.text_input(t['lbl_degree'])
                    year = st.text_input(t['lbl_year'])
                    if st.form_submit_button(t['btn_add']): add_education(inst, degree, year); st.rerun()

            elif section_logic == "Certifications":
                for i, cert in enumerate(st.session_state['resume_data']['certifications']):
                    st.text(f"• {cert['name']}")
                    if st.button(f"{t['btn_remove']} {i}", key=f"del_cert_{i}"): remove_certification(i); st.rerun()
                with st.form("add_cert_form"):
                    name = st.text_input(t['lbl_name'])
                    issuer = st.text_input(t['lbl_issuer'])
                    year = st.text_input(t['lbl_year'])
                    if st.form_submit_button(t['btn_add']): add_certification(name, issuer, year); st.rerun()

            elif section_logic == "Languages":
                for i, lang in enumerate(st.session_state['resume_data'].get('languages', [])):
                    with st.expander(f"{lang['name']}"):
                        st.write(f"{t['lbl_conv']}: {lang['conv']} | {t['lbl_comp']}: {lang['comp']} | {t['lbl_writ']}: {lang['writ']}")
                        if st.button(f"{t['btn_remove']} {i}", key=f"del_lang_{i}"): remove_language(i); st.rerun()
                
                with st.form("add_lang_form"):
                    name = st.text_input(t['lbl_language'])
                    levels = [t['lbl_level_basic'], t['lbl_level_inter'], t['lbl_level_adv'], t['lbl_level_fluent'], t['lbl_level_native']]
                    c1, c2, c3 = st.columns(3)
                    conv = c1.selectbox(t['lbl_conv'], levels)
                    comp = c2.selectbox(t['lbl_comp'], levels)
                    writ = c3.selectbox(t['lbl_writ'], levels)
                    if st.form_submit_button(t['btn_add']):
                        if name: add_language(name, conv, comp, writ); st.rerun()

            elif section_logic == "Awards":
                for i, aw in enumerate(st.session_state['resume_data'].get('awards', [])):
                    with st.expander(f"{aw['title']}"):
                        st.write(f"{aw['issuer']} - {aw['date']}")
                        if st.button(f"{t['btn_remove']} {i}", key=f"del_aw_{i}"): remove_award(i); st.rerun()
                
                with st.form("add_award_form"):
                    title = st.text_input(t['lbl_award_title'])
                    issuer = st.text_input(t['lbl_award_issuer'])
                    date_str = st.text_input(t['lbl_award_date'])
                    if st.form_submit_button(t['btn_add']):
                        if title: add_award(title, issuer, date_str); st.rerun()

            elif section_logic == "Volunteering":
                for i, vol in enumerate(st.session_state['resume_data'].get('volunteering', [])):
                    with st.expander(f"{vol['role']} - {vol['org']}"):
                        st.write(f"{vol['start']} -> {vol['end']} ({vol['category']})")
                        if st.button(f"{t['btn_remove']} {i}", key=f"del_vol_{i}"): remove_volunteering(i); st.rerun()
                
                with st.form("add_vol_form"):
                    role = st.text_input(t['lbl_vol_role'])
                    org = st.text_input(t['lbl_vol_org'])
                    c1, c2 = st.columns(2)
                    start = c1.text_input(t['lbl_start'])
                    end = c2.text_input(t['lbl_end'])
                    category = st.text_input(t['lbl_vol_cat'])
                    if st.form_submit_button(t['btn_add']):
                        if role: add_volunteering(role, org, start, end, category); st.rerun()

        with col_prev:
            st.markdown(f"### {t['preview_title']}")
            pdf_bytes = generate_pdf(st.session_state['resume_data'], scale_factor, lang_code)
            st.download_button(label=t['download_btn'], data=pdf_bytes, file_name=f"resume_{lang_code}.pdf", mime="application/pdf", type="primary")
            
            if Document:
                # Agora o DOCX usa o scale_factor
                docx_bytes = generate_docx(st.session_state['resume_data'], scale_factor, lang_code)
                if docx_bytes:
                    st.download_button(label="📥 Baixar / Download .docx (Word)", data=docx_bytes, file_name=f"resume_{lang_code}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                st.warning("⚠️ Biblioteca 'python-docx' não detectada. Instale com `pip install python-docx` para habilitar exportação Word.")

            
            # CSS Dinâmico (Preview)
            st.markdown(f"""
                <style>
                .resume-preview {{ font-family: 'Times New Roman', Times, serif; background-color: white; padding: {2 * scale_factor}rem; border-radius: 5px; box-shadow: 0 2px 5px rgba(0,0,0,0.1); color: #000; line-height: {1.3 * scale_factor}; }}
                .resume-header {{ text-align: center; padding-bottom: {5 * scale_factor}px; margin-bottom: {5 * scale_factor}px; }}
                .resume-name {{ color: #000; font-size: {2.5 * scale_factor}em; font-weight: bold; margin: 0; }}
                .resume-stack {{ font-family: Arial, Helvetica, sans-serif; font-size: {1.1 * scale_factor}em; color: #000; margin-top: 5px; margin-bottom: 5px; }}
                .resume-contact {{ font-size: {0.9 * scale_factor}em; color: #000; margin-top: 5px; }}
                .resume-section-title {{ color: #000; font-size: {1.2 * scale_factor}em; font-weight: bold; margin-top: {12 * scale_factor}px; margin-bottom: {2 * scale_factor}px; text-transform: uppercase; text-align: center; }}
                .resume-item {{ margin-bottom: {10 * scale_factor}px; }}
                .resume-item-header {{ font-family: 'Times New Roman', serif; font-weight: bold; font-size: {1.05 * scale_factor}em; color: #000; margin-bottom: 2px; }}
                .resume-item-sub {{ font-style: italic; color: #000; font-size: {0.95 * scale_factor}em; font-family: 'Times New Roman', serif; margin-bottom: 2px; }}
                .resume-description {{ margin-top: 2px; font-size: {0.95 * scale_factor}em; text-align: justify; color: #000; }}
                .resume-paragraph {{ margin-bottom: {3 * scale_factor}px; }}
                .resume-achievements {{ margin-top: {8 * scale_factor}px; font-weight: bold; margin-bottom: {2 * scale_factor}px; color: #000; }}
                </style>
            """, unsafe_allow_html=True)

            # HTML Preview (Resume)
            data = st.session_state['resume_data']
            stack_html = f'<div class="resume-stack">{data.get("stack", "")}</div>' if data.get("stack") else ""
            html_content = f"""<div class="resume-preview"><div class="resume-header"><h1 class="resume-name">{data['contact']['name']}</h1>{stack_html}<div class="resume-contact">📞 {data['contact']['phone']} | ✉️ {data['contact']['email']}<br>📍 {data['contact']['location']} | 🔗 {data['contact']['linkedin']}</div></div>"""
            if data['summary']:
                html_content += f"""<div class="resume-section"><div class="resume-section-title">{t['summary_header']}</div><div class="resume-description">"""
                for line in data['summary'].split('\n'):
                    if line.strip(): html_content += f"<div class='resume-paragraph'>{line}</div>"
                
                # Feitos / Achievements Preview
                if data.get('achievements'):
                     ach_label = "FEITOS" if lang_code == 'pt' else "KEY ACHIEVEMENTS"
                     html_content += f"<div class='resume-achievements'>{ach_label}</div>"
                     for line in data.get('achievements').split('\n'):
                         if line.strip(): html_content += f"<div class='resume-paragraph'>{line}</div>"

                html_content += "</div></div>"
            if data['skills']: html_content += f"""<div class="resume-section"><div class="resume-section-title">{t['skills_header']}</div><div class="resume-description">{', '.join(data['skills'])}</div></div>"""
            if data['experience']:
                html_content += f"""<div class="resume-section"><div class="resume-section-title">{t['experience_header']}</div>"""
                for exp in data['experience']:
                    html_content += f"""<div class="resume-item"><div class="resume-item-header">{exp['position']} - {exp['company']}</div><div class="resume-item-sub">{exp['start']} - {exp['end']}</div><div class="resume-description">"""
                    for line in exp['description'].split('\n'):
                        if line.strip(): html_content += f"<div class='resume-paragraph'>{line}</div>"
                    html_content += "</div></div>"
                html_content += "</div>"
            if data['education']:
                html_content += f"""<div class="resume-section"><div class="resume-section-title">{t['education_header']}</div>"""
                for edu in data['education']:
                    html_content += f"""<div class="resume-item"><div class="resume-item-header">{edu['degree']}</div><div class="resume-item-sub">{edu['institution']} - {t['lbl_year']}: {edu['year']}</div></div>"""
                html_content += "</div>"
            if data['certifications']:
                html_content += f"""<div class="resume-section"><div class="resume-section-title">{t['certifications_header']}</div><ul>"""
                for cert in data['certifications']: html_content += f"""<li><b>{cert['name']}</b> ({cert['issuer']}, {cert['year']})</li>"""
                html_content += "</ul></div>"
            
            # Languages HTML Preview
            if data.get('languages'):
                html_content += f"""<div class="resume-section"><div class="resume-section-title">{t['languages_header']}</div><ul>"""
                for lang in data['languages']:
                    html_content += f"""<li><b>{lang['name']}</b> - {t['lbl_conv']}: {lang['conv']} | {t['lbl_comp']}: {lang['comp']} | {t['lbl_writ']}: {lang['writ']}</li>"""
                html_content += "</ul></div>"

            # Awards HTML Preview
            if data.get('awards'):
                html_content += f"""<div class="resume-section"><div class="resume-section-title">{t['awards_header']}</div><ul>"""
                for aw in data['awards']:
                    html_content += f"""<li>{aw['title']} | {t['connector_offered_by']} {aw['issuer']} | {aw['date']}</li>"""
                html_content += "</ul></div>"

            # Volunteering HTML Preview
            if data.get('volunteering'):
                html_content += f"""<div class="resume-section"><div class="resume-section-title">{t['volunteering_header']}</div><ul>"""
                for vol in data['volunteering']:
                    html_content += f"""<li><b>{vol['role']}</b> | {vol['org']} | {vol['start']} -> {vol['end']} | {vol['category']}</li>"""
                html_content += "</ul></div>"

            html_content += "</div>"
            st.markdown(html_content, unsafe_allow_html=True)

    # === TAB 2: COVER LETTER BUILDER ===
    with tab2:
        col_cl_edit, col_cl_prev = st.columns([1, 1.2])
        
        with col_cl_edit:
            st.header("Cover Letter Builder")
            cl_data = st.session_state['cover_letter_data']
            
            with st.expander(t['cl_recipient_header'], expanded=True):
                cl_data['recipient']['manager'] = st.text_input(t['lbl_manager'], cl_data['recipient']['manager'])
                cl_data['recipient']['company'] = st.text_input(t['lbl_company_cl'], cl_data['recipient']['company'])
                cl_data['recipient']['address'] = st.text_input(t['lbl_address'], cl_data['recipient']['address'])
            
            with st.expander(t['cl_hook_header'], expanded=True):
                cl_data['opening']['greeting'] = st.text_input(t['lbl_greeting'], cl_data['opening']['greeting'])
                cl_data['opening']['hook'] = st.text_area(t['lbl_hook'], cl_data['opening']['hook'], height=100)
            
            with st.expander(t['cl_narrative_header']):
                cl_data['narrative'] = st.text_area(t['lbl_narrative'], cl_data['narrative'], height=150)
            
            with st.expander(t['cl_competencies_header']):
                cl_data['competencies']['star1'] = st.text_area(t['lbl_star1'], cl_data['competencies']['star1'], height=100)
                cl_data['competencies']['star2'] = st.text_area(t['lbl_star2'], cl_data['competencies']['star2'], height=100)
            
            with st.expander(t['cl_alignment_header']):
                cl_data['alignment']['research'] = st.text_area(t['lbl_alignment'], cl_data['alignment']['research'], height=100)
                cl_data['alignment']['differentiation'] = st.text_area(t['lbl_differentiation'], cl_data['alignment']['differentiation'], height=100)
            
            with st.expander(t['cl_closing_header']):
                cl_data['closing'] = st.text_area(t['lbl_closing'], cl_data['closing'], height=80)

        with col_cl_prev:
            st.markdown(f"### {t['preview_title']}")
            cl_pdf_bytes = generate_cl_pdf(st.session_state['resume_data'], cl_data, scale_factor, lang_code)
            st.download_button(label=t['download_cl_btn'], data=cl_pdf_bytes, file_name=f"cover_letter_{lang_code}.pdf", mime="application/pdf", type="primary")
            
            # HTML Preview (Cover Letter)
            cl_html = f"""
            <div class="resume-preview">
                <div class="resume-header">
                    <h1 class="resume-name">{data['contact']['name']}</h1>
                    <div class="resume-contact">📞 {data['contact']['phone']} | ✉️ {data['contact']['email']}</div>
                </div>
                <div style="margin-top: 20px; font-family: Arial; line-height: 1.6;">
                    <p><strong>{date.today().strftime('%B %d, %Y')}</strong></p>
                    <p>{cl_data['recipient']['manager']}<br>{cl_data['recipient']['company']}<br>{cl_data['recipient']['address']}</p>
                    <p>{cl_data['opening']['greeting']}:</p>
                    <p>{cl_data['opening']['hook']}</p>
                    <p>{cl_data['narrative']}</p>
                    <p>{cl_data['competencies']['star1']}</p>
                    <p>{cl_data['competencies']['star2']}</p>
                    <p>{cl_data['alignment']['research']}</p>
                    <ul>
            """
            for line in cl_data['alignment']['differentiation'].split('\n'):
                if line.strip(): cl_html += f"<li>{line}</li>"
            
            cl_html += f"""
                    </ul>
                    <p>{cl_data['closing']}</p>
                    <br>
                    <p>Sincerely,</p>
                    <br>
                    <p><strong>{data['contact']['name']}</strong></p>
                </div>
            </div>
            """
            st.markdown(cl_html, unsafe_allow_html=True)

    # === TAB 3: PROPOSAL BUILDER (ABNT) ===
    with tab3:
        st.header(t.get('prop_header', 'Proposta'))
        p_data = st.session_state['proposal_data']

        # 1. CAPA & FOLHA DE ROSTO
        with st.expander("📘 Capa & Folha de Rosto (ABNT)", expanded=True):
            c1, c2 = st.columns(2)
            p_data['cover']['author'] = c1.text_input(t['lbl_author_prop'], p_data['cover']['author'])
            p_data['cover']['institution'] = c2.text_input(t['lbl_institution'], p_data['cover']['institution'])
            p_data['cover']['title'] = st.text_input(t['lbl_project_title'], p_data['cover']['title'])
            p_data['cover']['subtitle'] = st.text_input(t['lbl_subtitle'], p_data['cover']['subtitle'])
            c3, c4 = st.columns(2)
            p_data['cover']['city'] = c3.text_input(t['lbl_city'], p_data['cover']['city'])
            p_data['cover']['year'] = c4.text_input(t['lbl_year'], p_data['cover']['year'])
            p_data['cover']['theme'] = st.text_area(t['lbl_theme'], p_data['cover']['theme'], height=70)

        # 2. SEÇÕES DO PROJETO (1-9)
        st.divider()
        st.subheader("Estrutura do Projeto")
        
        sections_map = [
            ("1", t['sec_1']), ("2", t['sec_2']), ("3", t['sec_3']), ("4", t['sec_4']),
            ("5", t['sec_5']), ("6", t['sec_6']), ("7", t['sec_7']), ("8", t['sec_8']), ("9", t['sec_9'])
        ]

        for sec_id, sec_title in sections_map:
            with st.expander(f"{sec_id}. {sec_title}"):
                
                # Lógica Especial: Cronograma (5)
                if sec_id == "5":
                    for i, item in enumerate(p_data['timeline']):
                        c1, c2, c3 = st.columns([0.2, 0.7, 0.1])
                        c1.text(item['date'])
                        c2.text(item['milestone'])
                        if c3.button("X", key=f"del_time_{i}"): remove_timeline_milestone(i); st.rerun()
                    with st.form("add_timeline_form"):
                        c1, c2 = st.columns([0.3, 0.7])
                        d_date = c1.date_input(t.get('lbl_date', 'Data'))
                        d_milestone = c2.text_input(t.get('lbl_milestone', 'Marco'))
                        if st.form_submit_button(t['btn_add']):
                            if d_milestone: add_timeline_milestone(d_date, d_milestone); st.rerun()

                # Lógica Especial: Orçamento (6)
                elif sec_id == "6":
                    total = sum(item['amount'] for item in p_data['budget'])
                    st.markdown(f"**Total: {total:,.2f}**")
                    for i, item in enumerate(p_data['budget']):
                        c1, c2, c3 = st.columns([0.6, 0.3, 0.1])
                        c1.text(item['item'])
                        c2.text(f"{item['amount']:.2f}")
                        if c3.button("X", key=f"del_budget_{i}"): remove_budget_item(i); st.rerun()
                    with st.form("add_budget_form"):
                        c1, c2 = st.columns([0.7, 0.3])
                        b_item = c1.text_input(t.get('lbl_budget_item', 'Item'))
                        b_amount = c2.number_input(t.get('lbl_amount', 'Valor'), min_value=0.0, step=100.0)
                        if st.form_submit_button(t['btn_add']):
                            if b_item: add_budget_item(b_item, b_amount); st.rerun()

                # Lógica Padrão: Texto com Subseções
                else:
                    subsections = p_data['text_sections'].get(sec_id, [])
                    for i, sub in enumerate(subsections):
                        st.markdown(f"**{sec_id}.{i+1} {sub['title']}**")
                        st.text(sub['content'][:60] + "..." if len(sub['content']) > 60 else sub['content'])
                        if st.button(f"{t['btn_remove']} {i}", key=f"del_sub_{sec_id}_{i}"): remove_text_subsection(sec_id, i); st.rerun()
                    
                    with st.form(f"add_sub_{sec_id}"):
                        s_title = st.text_input(t['lbl_sub_title'])
                        s_content = st.text_area(t['lbl_sub_content'], height=100)
                        if st.form_submit_button(t['btn_add']):
                            if s_title and s_content: add_text_subsection(sec_id, s_title, s_content); st.rerun()

        # Download
        st.divider()
        prop_pdf_bytes = generate_proposal_pdf(p_data, t)
        st.download_button(label=t.get('download_prop_btn', 'Baixar PDF'), data=prop_pdf_bytes, file_name=f"projeto_{lang_code}.pdf", mime="application/pdf", type="primary")

    # === TAB 4: REPORT BUILDER (ABNT) ===
    with tab4:
        st.header(t.get('rep_header', 'Relatório Técnico'))
        r_data = st.session_state['report_data']

        # 1. CAPA & FOLHA DE ROSTO
        with st.expander("📘 Capa & Folha de Rosto (ABNT)", expanded=True):
            c1, c2 = st.columns(2)
            r_data['cover']['author'] = c1.text_input(t['lbl_author_prop'], r_data['cover']['author'], key="rep_auth")
            r_data['cover']['institution'] = c2.text_input(t['lbl_institution'], r_data['cover']['institution'], key="rep_inst")
            r_data['cover']['title'] = st.text_input(t['lbl_project_title'], r_data['cover']['title'], key="rep_title")
            r_data['cover']['subtitle'] = st.text_input(t['lbl_subtitle'], r_data['cover']['subtitle'], key="rep_sub")
            c3, c4 = st.columns(2)
            r_data['cover']['city'] = c3.text_input(t['lbl_city'], r_data['cover']['city'], key="rep_city")
            r_data['cover']['year'] = c4.text_input(t['lbl_year'], r_data['cover']['year'], key="rep_year")
            r_data['cover']['theme'] = st.text_area(t['lbl_theme'], r_data['cover']['theme'], height=70, key="rep_theme")

        # 2. SEÇÕES DO RELATÓRIO (1-7)
        st.divider()
        st.subheader("Conteúdo do Relatório")
        
        rep_sections = [
            ("1", t['rep_sec_1']), ("2", t['rep_sec_2']), ("3", t['rep_sec_3']),
            ("4", t['rep_sec_4']), ("5", t['rep_sec_5']), ("6", t['rep_sec_6']),
            ("7", t['rep_sec_7'])
        ]

        for sec_id, sec_title in rep_sections:
            with st.expander(f"{sec_id}. {sec_title}"):
                subsections = r_data['text_sections'].get(sec_id, [])
                for i, sub in enumerate(subsections):
                    st.markdown(f"**{sec_id}.{i+1} {sub['title']}**")
                    st.text(sub['content'][:60] + "..." if len(sub['content']) > 60 else sub['content'])
                    if st.button(f"{t['btn_remove']} {i}", key=f"del_rep_sub_{sec_id}_{i}"): remove_report_subsection(sec_id, i); st.rerun()
                
                with st.form(f"add_rep_sub_{sec_id}"):
                    s_title = st.text_input(t['lbl_sub_title'])
                    s_content = st.text_area(t['lbl_sub_content'], height=100)
                    if st.form_submit_button(t['btn_add']):
                        if s_title and s_content: add_report_subsection(sec_id, s_title, s_content); st.rerun()

        # 3. ANEXOS (IMAGENS)
        st.divider()
        with st.expander(f"8. {t['rep_sec_8']} (Imagens)"):
            for i, annex in enumerate(r_data['annexes']):
                st.image(annex['image'], caption=f"Fig {i+1}: {annex['caption']}", width=150)
                if st.button(f"{t['btn_remove']} {i}", key=f"del_annex_{i}"): remove_report_annex(i); st.rerun()
            
            with st.form("add_annex_form"):
                uploaded_file = st.file_uploader(t['lbl_img_upload'], type=['png', 'jpg', 'jpeg'])
                caption = st.text_input(t['lbl_img_caption'])
                if st.form_submit_button(t['btn_add']):
                    if uploaded_file and caption:
                        add_report_annex(uploaded_file.getvalue(), caption)
                        st.rerun()

        # Download
        st.divider()
        rep_pdf_bytes = generate_report_pdf(r_data, t)
        st.download_button(label=t.get('download_rep_btn', 'Baixar Relatório PDF'), data=rep_pdf_bytes, file_name=f"relatorio_{lang_code}.pdf", mime="application/pdf", type="primary")

    
    # === TAB 5: ATS SIMULATOR ===
    with tab5:
        st.header(t['ats_header'])
        st.markdown(t['ats_desc'])
        
        if PdfReader is None:
            st.error(t['ats_error_lib'])
        else:
            uploaded_pdf = st.file_uploader(t['ats_upload'], type=["pdf"])
            
            if uploaded_pdf is not None:
                if st.button("🔍 Analisar PDF"):
                    with st.spinner("Processando como um robô (ATS)..."):
                        # 1. Extração
                        raw_text = extract_text_from_pdf(uploaded_pdf)
                        
                        # 2. Análise
                        score, parsed_data, sections_found = analyze_ats_compatibility(raw_text)
                        
                        # 3. Atualizar Estado
                        st.session_state['ats_data']['score'] = score
                        st.session_state['ats_data']['raw_text'] = raw_text
                        st.session_state['ats_data']['sections_found'] = sections_found
                        st.session_state['ats_data']['parsed_content'] = parsed_data
            
            # Dashboard de Resultados
            if st.session_state['ats_data']['raw_text']:
                st.divider()
                
                # Métricas de Topo
                c1, c2, c3 = st.columns(3)
                final_score = st.session_state['ats_data']['score']
                score_color = "green" if final_score > 80 else "orange" if final_score > 50 else "red"
                
                c1.metric(label=t['ats_score'], value=f"{final_score}/100")
                c2.metric(label=t['ats_text_len'], value=len(st.session_state['ats_data']['raw_text']))
                c3.metric(label=t['ats_sec_found'], value=f"{len(st.session_state['ats_data']['sections_found'])}/7")
                
                # Visualização de Dados
                col_left, col_right = st.columns(2)
                
                with col_left:
                    st.subheader(t['ats_parsed_data'])
                    st.json(st.session_state['ats_data']['parsed_content'])
                    
                with col_right:
                    st.subheader(t['ats_raw_text'])
                    st.text_area("Raw Output", st.session_state['ats_data']['raw_text'], height=400, disabled=True)

if __name__ == "__main__":
    main()
